from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REFERENCE = Path(
    "/Users/derekdreibrodt/Downloads/"
    "WD_Web_Solutions_Web_Development_Agreement_TEMPLATE.docx"
)
OUTPUT = Path(
    "/Users/derekdreibrodt/Documents/repos/austinsurfacepros.com/docs/contracts/"
    "WD_Web_Solutions_Austin_Surface_Pros_Web_Development_Agreement_REVIEW.docx"
)

FONT = "Times New Roman"
BODY_SIZE = Pt(11)
HIGHLIGHT = WD_COLOR_INDEX.YELLOW


def set_run_font(run, size=BODY_SIZE, bold=None, italic=None) -> None:
    run.font.name = FONT
    run.font.size = size
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_text(paragraph, text: str, *, changed=True, bold=False, italic=False, size=BODY_SIZE):
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    if changed:
        run.font.highlight_color = HIGHLIGHT
    return run


def format_body(paragraph, *, after=Pt(7), before=Pt(0), keep=False) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = before
    paragraph.paragraph_format.space_after = after
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.keep_together = keep


def add_body(doc, text: str, *, changed=True, bold=False, italic=False, after=Pt(7), keep=False):
    paragraph = doc.add_paragraph()
    format_body(paragraph, after=after, keep=keep)
    add_text(paragraph, text, changed=changed, bold=bold, italic=italic)
    return paragraph


def add_mixed(doc, pieces, *, after=Pt(7), keep=False):
    paragraph = doc.add_paragraph()
    format_body(paragraph, after=after, keep=keep)
    for piece in pieces:
        if isinstance(piece, str):
            add_text(paragraph, piece)
        else:
            text = piece[0]
            options = piece[1]
            add_text(paragraph, text, **options)
    return paragraph


def add_heading(doc, text: str, *, changed=True):
    source_heading = next(style for style in doc.styles if style.style_id == "Heading1")
    paragraph = doc.add_paragraph(style=source_heading)
    paragraph.paragraph_format.keep_with_next = True
    add_text(paragraph, text.upper(), changed=changed, bold=True)
    return paragraph


def apply_numpr(paragraph, numpr) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    current = ppr.find(qn("w:numPr"))
    if current is not None:
        ppr.remove(current)
    ppr.append(deepcopy(numpr))


def add_bullet(doc, text: str, numpr, *, level=0, changed=True, bold=False):
    paragraph = doc.add_paragraph(style="List Paragraph")
    apply_numpr(paragraph, numpr[level])
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.keep_together = True
    add_text(paragraph, text, changed=changed, bold=bold)
    return paragraph


def add_page_break(doc) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_signature_line(doc, label: str, *, changed=True) -> None:
    paragraph = doc.add_paragraph()
    format_body(paragraph, after=Pt(12))
    add_text(
        paragraph,
        f"{label}: ____________________________________________     Date: __________________",
        changed=changed,
    )


def set_cellless_tab_stop(paragraph, position_inches: float) -> None:
    tabs = paragraph.paragraph_format.tab_stops
    tabs.add_tab_stop(Inches(position_inches))


def add_section_numbering(paragraph, number: str, title: str) -> None:
    add_text(paragraph, f"{number}. {title.upper()}", changed=True, bold=True)


def clear_document_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_document_language(doc: Document) -> None:
    styles = doc.styles.element
    doc_defaults = styles.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        styles.insert(0, doc_defaults)
    rpr_default = doc_defaults.find(qn("w:rPrDefault"))
    if rpr_default is None:
        rpr_default = OxmlElement("w:rPrDefault")
        doc_defaults.append(rpr_default)
    rpr = rpr_default.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        rpr_default.append(rpr)
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    lang.set(qn("w:val"), "en-US")


def build() -> None:
    doc = Document(str(REFERENCE))

    numpr = {}
    for paragraph in doc.paragraphs:
        ppr = paragraph._p.pPr
        if ppr is None or ppr.numPr is None or ppr.numPr.ilvl is None:
            continue
        level = int(ppr.numPr.ilvl.val)
        if level in (0, 1) and level not in numpr:
            numpr[level] = deepcopy(ppr.numPr)
    if set(numpr) != {0, 1}:
        raise RuntimeError("Could not find both source bullet levels")

    clear_document_body(doc)
    set_document_language(doc)

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    if "Normal" in [style.name for style in doc.styles]:
        normal = doc.styles["Normal"]
    else:
        normal = doc.styles.add_style("Normal", WD_STYLE_TYPE.PARAGRAPH)
    normal.font.name = FONT
    normal.font.size = BODY_SIZE
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), FONT)
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), FONT)

    heading_style = next(style for style in doc.styles if style.style_id == "Heading1")
    heading_style.base_style = normal
    heading_style.font.name = FONT
    heading_style.font.size = BODY_SIZE
    heading_style.font.bold = True
    heading_style.font.color.rgb = RGBColor(0, 0, 0)
    heading_style.paragraph_format.space_before = Pt(9)
    heading_style.paragraph_format.space_after = Pt(3)
    heading_style.paragraph_format.keep_with_next = True
    heading_style.paragraph_format.keep_together = True

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    add_text(title, "Web Development Agreement", changed=False, bold=True, size=Pt(14))

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(14)
    add_text(subtitle, "Austin Surface Pros", changed=True, bold=True)

    add_mixed(
        doc,
        [
            "This Web Development, Managed Hosting, and Support Agreement (the \u201cAgreement\u201d) is effective on the date of the last signature below (the \u201cEffective Date\u201d) and is between ",
            ("Wayne Filer d/b/a Austin Surface Pros", {"changed": True, "bold": True}),
            (" (\u201cClient\u201d), with a notice and billing address of 1239 County Road 250, Georgetown, Texas 78633 and email address austinsurfacepros@gmail.com, and ", {"changed": True}),
            ("WD Web Solutions LLC", {"changed": True, "bold": True}),
            (" (\u201cWD\u201d), of 324 Highland Springs Lane, Georgetown, Texas 78633. Client and WD may each be a \u201cParty\u201d and together the \u201cParties.\u201d", {"changed": True}),
        ],
    )

    add_body(
        doc,
        "WHEREAS, WD provides website design, development, cloud hosting, domain management, and related support services; and",
        changed=True,
    )
    add_body(
        doc,
        "WHEREAS, Client wishes to engage WD to design, launch, host, and support a business website for Austin Surface Pros on the terms below.",
        changed=True,
    )
    add_body(
        doc,
        "NOW, THEREFORE, in consideration of the mutual promises in this Agreement, the Parties agree as follows:",
        changed=True,
    )

    add_heading(doc, "1. Engagement and Contract Documents")
    add_body(
        doc,
        "Client retains WD to perform the website design, development, launch, managed hosting, domain management, maintenance, and support services described in this Agreement. Exhibit A (Project Scope) and Exhibit B (Managed Hosting and Support Schedule) are incorporated into this Agreement. If there is a conflict, the main body of this Agreement controls, followed by Exhibit B and then Exhibit A.",
    )

    add_heading(doc, "2. Project Schedule, Review, and Acceptance")
    add_body(
        doc,
        "WD will begin work after the Effective Date and payment of the Initial Fee. The target production launch date is September 1, 2026. This is a target, not a guarantee, and assumes Client provides timely content, access, approvals, and consolidated feedback. Client-caused delays automatically extend the schedule by at least the period of delay and any reasonable remobilization time.",
    )
    add_body(
        doc,
        "The Initial Fee includes two consolidated rounds of design feedback. Client will use reasonable efforts to provide one consolidated written response within five business days after each review request. Fragmented feedback, additional review rounds, or changes outside the accepted scope may be billed at the Additional Services Rate after notice and approval under Section 5.",
    )
    add_body(
        doc,
        "The demo website may contain seed data, sample posts, placeholder text, stock-style imagery, or other illustrative content. Such material is for review only and is not a representation of Client's actual services, credentials, prices, or business claims. Client must approve or replace it before launch.",
    )
    add_body(
        doc,
        "The website will be accepted upon the earliest of: (a) Client's written approval; (b) Client's authorization to publish; (c) Client's production use of the website; or (d) five business days after final delivery without written notice identifying a material failure to meet Exhibit A. WD will correct a timely reported material scope defect as Client's exclusive acceptance remedy.",
    )

    add_heading(doc, "3. Client Responsibilities and Dependencies")
    add_body(doc, "Client will, at its own expense:")
    for item in [
        "Provide accurate final copy, service descriptions, business information, pricing, logos, photographs, videos, and other media in usable formats, together with all approvals needed for WD to use them.",
        "Provide timely decisions, consolidated feedback, and access to the domain registrar, DNS, Google Business Profile, YouTube or other media sources, and any other systems reasonably required for the work.",
        "Provide necessary authentication, usernames, permissions, and temporary credentials through a reasonably secure method. Client will promptly rotate any shared password after access is established. WD is not responsible for delay caused by unavailable, incomplete, or invalid access.",
        "Ensure that Client Materials, business claims, testimonials, offers, and instructions are accurate, lawful, noninfringing, and properly licensed, and obtain any releases required for people, properties, logos, music, photographs, or video shown on the website.",
        "Identify the individuals authorized to receive separate administrator accounts. Public registration, public customer accounts, public comments, and topic subscriptions are excluded from the scope.",
    ]:
        add_bullet(doc, item, numpr)

    add_heading(doc, "4. Fees and Payment")
    add_body(
        doc,
        "Initial Fee. Client will pay a one-time design and setup fee of $300.00 upon signing this Agreement in August 2026. WD is not required to begin or continue work until the Initial Fee is paid.",
    )
    add_body(
        doc,
        "Monthly Fee. Client will pay $50.00 per month in advance beginning September 1, 2026, due on the first day of each month. September billing will begin as scheduled if the website is production-ready or launch is delayed by Client's missing content, access, feedback, or approval. If WD alone causes a material delay that prevents production readiness, monthly billing will begin when the website is production-ready.",
    )
    add_body(
        doc,
        "Invoices may be delivered and paid electronically, including through Square. Amounts not disputed in good faith are due as stated on the invoice. After written notice, WD may suspend work, hosting, or access if an invoice remains unpaid for ten calendar days after its due date. Suspension does not waive amounts already due or Client's responsibility for costs incurred during suspension.",
    )
    add_body(
        doc,
        "Unless an invoice states otherwise, fees are exclusive of applicable sales or similar taxes. Client is responsible for taxes legally chargeable to Client, but not taxes based on WD's net income.",
    )

    add_heading(doc, "5. Included Changes, Meetings, and Additional Services")
    add_body(
        doc,
        "The Monthly Fee includes up to five minor content edits totaling no more than 30 minutes of WD labor per calendar month. A minor content edit is a small change to existing text, an existing image, business contact information, or a comparable item that does not alter the website's architecture, design system, data model, integration, workflow, or legal obligations. Unused edits or time do not roll over.",
    )
    add_body(
        doc,
        "The Monthly Fee also includes ad hoc meetings scheduled by mutual availability totaling up to one hour per calendar month. Unused meeting time does not roll over. WD may waive a meeting charge in a particular instance without creating a continuing right to free time.",
    )
    add_body(
        doc,
        "Work beyond the included allowances is billed at $75.00 per hour in 15-minute increments (the \u201cAdditional Services Rate\u201d). WD will warn Client when WD reasonably expects requested changes or meetings to exceed an allowance and will obtain written approval before performing billable additional work. Email or an electronic approval through the invoicing or signature platform is sufficient. Substantial changes, new features, integrations, redesigns, or work that materially changes scope require a signed change order or separate agreement before work begins.",
    )

    add_heading(doc, "6. Twelve-Month Price Protection and Renewal")
    add_body(
        doc,
        "The $50.00 Monthly Fee, $75.00 Additional Services Rate, included edit and meeting allowances, and the operating-cost allowance and handling margin in Exhibit B are fixed through August 31, 2027. Pass-through third-party costs may change as their providers change prices and remain subject to Exhibit B.",
    )
    add_body(
        doc,
        "WD may propose different pricing, allowances, or service terms for any renewal by giving at least 60 days' written notice before August 31, 2027. No proposed increase is effective unless Client approves it in a written renewal or amendment signed electronically or on paper. If the Parties do not approve a renewal, neither Party is required to continue services after August 31, 2027, and Section 16 will govern transition and handoff.",
    )

    add_heading(doc, "7. Third-Party Services and Excluded Costs")
    add_body(
        doc,
        "AWS, domain registries, certificate authorities, email providers, video platforms, Google, Square, and other third parties supply services under their own terms and may change or discontinue them. WD is not responsible for third-party outages, price changes, policy changes, data loss, suspension, or acts or omissions beyond WD's reasonable control.",
    )
    add_body(
        doc,
        "A business-domain mailbox service such as Google Workspace, Microsoft 365, or a comparable provider is optional, is not included in the Monthly Fee or operating-cost allowance, and will be paid directly by Client unless the Parties sign a written amendment. Transactional email used by the contact or estimate form is included only to the extent covered by Exhibit B's operating-cost allowance.",
    )

    add_heading(doc, "8. Domain and Cloud Account Administration")
    add_body(
        doc,
        "Client is the beneficial owner of the austinsurfacepros.com domain and authorizes WD to transfer, register, renew, configure, and manage the domain, DNS, certificates, and website resources within an AWS account or AWS Organization controlled by WD. Administrative control by WD does not transfer Client's beneficial ownership of the domain.",
    )
    add_body(
        doc,
        "WD will use commercially reasonable practices to segregate Client resources and maintain account records. Client will keep its ownership and contact information current. Upon termination, full payment, and receipt of the destination registrar or AWS account information, WD will provide an ordinary domain transfer authorization or initiate an ordinary transfer within ten business days. One ordinary domain transfer is included. Unusually complex migration, data transformation, third-party troubleshooting, or post-handoff assistance is billed at the Additional Services Rate after approval.",
    )

    add_heading(doc, "9. Intellectual Property and Licenses")
    add_body(
        doc,
        "Client retains ownership of the domain and all text, logos, trademarks, photographs, videos, data, and other materials supplied by or for Client (\u201cClient Materials\u201d). Client grants WD a nonexclusive license to host, copy, modify, display, transmit, and otherwise use Client Materials solely to perform the Services.",
    )
    add_body(
        doc,
        "WD retains ownership of all pre-existing or independently developed tools, reusable components, libraries, templates, deployment systems, infrastructure code, generalized code, processes, designs, know-how, and improvements (\u201cWD Materials\u201d). Subject to full payment, WD grants Client a nonexclusive, perpetual, worldwide license to use the WD Materials only as incorporated into the delivered website for Client's business. Client may authorize a successor provider to operate and maintain that website for Client, but may not sell, sublicense, extract, or reuse WD Materials as a stand-alone product or service. Open-source and third-party materials remain subject to their applicable licenses.",
    )

    add_heading(doc, "10. Privacy, Accessibility, and Legal Compliance")
    add_body(
        doc,
        "WD will implement the technical controls, collection notices, privacy settings, and accessibility features expressly included in Exhibit A using commercially reasonable care. Client is responsible for reviewing and approving Client Materials, business claims, service-area statements, privacy disclosures, terms, consent language, retention requirements, and other legal requirements applicable to Client's business and use of the website.",
    )
    add_body(
        doc,
        "WD does not provide legal, tax, accounting, or regulatory advice and does not guarantee that the website or sample policies satisfy every law, regulation, industry standard, or accessibility requirement. Client should obtain qualified legal review before launch and whenever its business practices or website data flows change. Neither Party will knowingly instruct the other to violate applicable law.",
    )

    add_heading(doc, "11. Confidentiality and Security")
    add_body(
        doc,
        "Each Party will use the other Party's nonpublic business, technical, credential, and customer information only to perform or receive the Services and will protect it using reasonable care. This duty does not apply to information that is public through no breach, already lawfully known, independently developed, or lawfully obtained from another source. A Party may disclose information when legally required after giving notice when permitted. These obligations survive termination for three years; trade-secret obligations survive while the information remains a trade secret.",
    )
    add_body(
        doc,
        "WD will use commercially reasonable administrative and technical safeguards appropriate to the Services, but no Internet service, credential-sharing method, backup, or security control is error-free. Client will use unique administrator credentials, protect its accounts, promptly report suspected compromise, and avoid sending passwords by ordinary unencrypted email.",
    )

    add_heading(doc, "12. Limited Warranty and Disclaimers")
    add_body(
        doc,
        "WD warrants that it will perform the Services in a professional and workmanlike manner. Client must notify WD in writing of a claimed breach within 30 days after the affected Service is delivered. WD's exclusive obligation, and Client's exclusive remedy, is for WD to reperform or correct the materially nonconforming Service when reasonably possible.",
    )
    add_body(
        doc,
        "EXCEPT FOR THE EXPRESS WARRANTY ABOVE, THE SERVICES, WEBSITE, HOSTING, SUPPORT, AND ALL WD MATERIALS ARE PROVIDED \u201cAS IS\u201d AND \u201cAS AVAILABLE.\u201d TO THE MAXIMUM EXTENT PERMITTED BY LAW, WD DISCLAIMS IMPLIED WARRANTIES OF MERCHANTABILITY, FITNESS FOR A PARTICULAR PURPOSE, TITLE, NON-INFRINGEMENT, AND ANY WARRANTY ARISING FROM COURSE OF DEALING OR USAGE. WD DOES NOT GUARANTEE SEARCH RANKING, TRAFFIC, LEADS, SALES, CONTINUOUS AVAILABILITY, ERROR-FREE OPERATION, DATA RECOVERY, SECURITY AGAINST EVERY THREAT, OR RESULTS FROM THIRD-PARTY SERVICES.",
    )

    add_heading(doc, "13. Limitation of Liability")
    add_body(
        doc,
        "TO THE MAXIMUM EXTENT PERMITTED BY LAW, NEITHER PARTY WILL BE LIABLE FOR INDIRECT, INCIDENTAL, SPECIAL, EXEMPLARY, PUNITIVE, OR CONSEQUENTIAL DAMAGES, OR FOR LOST PROFITS, LOST REVENUE, LOSS OF GOODWILL, OR LOSS OR CORRUPTION OF DATA, EVEN IF ADVISED OF THE POSSIBILITY.",
    )
    add_body(
        doc,
        "EXCEPT FOR CLIENT'S PAYMENT OBLIGATIONS, A PARTY'S FRAUD, GROSS NEGLIGENCE OR WILLFUL MISCONDUCT, OR A PARTY'S EXPRESS INDEMNIFICATION OBLIGATIONS, EACH PARTY'S TOTAL AGGREGATE LIABILITY ARISING OUT OF THIS AGREEMENT WILL NOT EXCEED THE FEES PAID OR PAYABLE TO WD FOR SERVICES DURING THE 12 MONTHS PRECEDING THE EVENT GIVING RISE TO THE CLAIM, EXCLUDING TAXES AND PASS-THROUGH THIRD-PARTY COSTS. The limitations apply regardless of the legal theory and are an essential basis of the Parties' bargain.",
    )

    add_heading(doc, "14. Mutual Indemnification")
    add_body(
        doc,
        "Client will defend, indemnify, and hold harmless WD and its personnel from third-party claims, damages, and reasonable costs arising from Client Materials, Client's business claims or services, Client's unlawful instructions, or Client's material breach of this Agreement, except to the extent caused by WD's gross negligence or willful misconduct.",
    )
    add_body(
        doc,
        "WD will defend, indemnify, and hold harmless Client from third-party claims that original deliverables created solely by WD under this Agreement directly infringe a United States copyright or trademark, or that arise from WD's gross negligence or willful misconduct, except to the extent caused by Client Materials, Client instructions, modifications not made by WD, or use outside the Agreement. WD may modify or replace an allegedly infringing item or terminate the affected Service and refund prepaid fees for the unused affected period.",
    )
    add_body(
        doc,
        "The indemnified Party must promptly notify the indemnifying Party, provide reasonable cooperation at the indemnifying Party's expense, and allow the indemnifying Party to control the defense and settlement. No settlement may admit fault by or impose a nonmonetary obligation on the indemnified Party without its written consent.",
    )

    add_heading(doc, "15. Term and Termination")
    add_body(
        doc,
        "The initial service term begins on the Effective Date and ends August 31, 2027 unless terminated earlier. Either Party may terminate for convenience on 60 days' written notice. Either Party may terminate for a material breach not cured within ten calendar days after written notice. WD may suspend or terminate immediately to address illegal activity, a material security risk, or persistent nonpayment after the notice described in Section 4.",
    )
    add_body(
        doc,
        "Termination does not excuse payment of fees, approved Additional Services, or third-party costs incurred through the effective termination date. Prepaid Monthly Fees are nonrefundable except when WD terminates for convenience without Client breach, in which case WD will refund the unused full-month portion after the effective termination date.",
    )

    add_heading(doc, "16. Transition After Termination")
    add_body(
        doc,
        "After termination and full payment, WD will provide Client with a reasonable export of Client Materials and initiate the ordinary domain handoff described in Section 8. WD may retain archival copies as required by law or routine backup practices, subject to confidentiality obligations. Unless otherwise agreed, WD may delete hosted website data 30 days after completing the handoff. Client is responsible for arranging replacement hosting before that date. Any extended hosting, custom export, migration, or successor-provider support is billed at the Additional Services Rate after approval.",
    )

    add_heading(doc, "17. Relationship, Assignment, and Personnel")
    add_body(
        doc,
        "WD is an independent contractor and not Client's employee, partner, joint venturer, fiduciary, or agent. Neither Party may bind the other. WD may use employees and subcontractors and remains responsible for their performance. Neither Party may assign this Agreement without the other's written consent, except to a successor in connection with a merger, reorganization, or sale of substantially all relevant assets, provided the successor assumes the assigning Party's obligations.",
    )

    add_heading(doc, "18. Force Majeure")
    add_body(
        doc,
        "Neither Party is liable for delay or failure caused by events beyond its reasonable control, including severe weather, fire, natural disaster, labor disruption, war, terrorism, epidemic, governmental action, utility or Internet failure, cyberattack not caused by that Party's failure to use reasonable safeguards, or failure of AWS or another third-party provider. Payment obligations for Services already performed are not excused.",
    )

    add_heading(doc, "19. Notices")
    add_body(
        doc,
        "Notices under this Agreement must be in writing and delivered by personal delivery, nationally recognized overnight courier, certified U.S. mail with return receipt, or email to the physical or email address stated above or later designated in writing. Email notice is effective when sent without an automated failure notice; all other notice is effective upon documented delivery. Routine service requests, approvals, and scheduling may be handled by email or the Parties' agreed electronic platform.",
    )

    add_heading(doc, "20. General Terms")
    add_body(
        doc,
        "This Agreement and its exhibits are the entire agreement concerning their subject and supersede prior or contemporaneous proposals and discussions. A purchase order or invoice term does not modify this Agreement. Any amendment, waiver, renewal, or change order must be in writing and accepted by both Parties. A waiver in one instance is not a continuing waiver.",
    )
    add_body(
        doc,
        "If a provision is held invalid or unenforceable, it will be limited to the minimum extent necessary and the remaining provisions will continue in effect. Headings are for convenience only. Sections that by their nature should survive termination will survive, including payment, ownership, confidentiality, disclaimers, limitations, indemnification, and transition obligations.",
    )
    add_body(
        doc,
        "Texas law governs this Agreement without regard to conflict-of-law rules. The state and federal courts located in Williamson County, Texas have exclusive jurisdiction, and each Party consents to that venue.",
    )
    add_body(
        doc,
        "This Agreement may be signed in counterparts and through DocuSign, Square, or another electronic signature platform. Electronic signatures, electronic records, and electronically delivered copies are intended to have the same effect as originals.",
    )

    add_page_break(doc)
    add_heading(doc, "Signatures")
    add_body(
        doc,
        "The undersigned agree to be bound by this Agreement, including Exhibit A and Exhibit B, as of the Effective Date.",
    )

    client_label = doc.add_paragraph()
    client_label.paragraph_format.space_before = Pt(12)
    client_label.paragraph_format.space_after = Pt(4)
    add_text(client_label, "CLIENT:", changed=False, bold=True)
    add_body(doc, "Wayne Filer d/b/a Austin Surface Pros", bold=True, after=Pt(12))
    add_signature_line(doc, "By")
    add_body(doc, "Wayne Filer", bold=True, after=Pt(3))
    add_body(doc, "Authorized Signer", after=Pt(18))

    wd_label = doc.add_paragraph()
    wd_label.paragraph_format.space_before = Pt(12)
    wd_label.paragraph_format.space_after = Pt(4)
    add_text(wd_label, "WEB DEVELOPER:", changed=False, bold=True)
    add_body(doc, "WD Web Solutions LLC", changed=False, after=Pt(12))
    add_signature_line(doc, "By", changed=False)
    add_body(doc, "Zachary Whittenton", changed=False, after=Pt(3))
    add_body(doc, "Member", changed=False)

    add_page_break(doc)
    exhibit_a = doc.add_paragraph()
    exhibit_a.alignment = WD_ALIGN_PARAGRAPH.CENTER
    exhibit_a.paragraph_format.space_after = Pt(3)
    add_text(exhibit_a, "EXHIBIT A", bold=True)
    exhibit_a_title = doc.add_paragraph()
    exhibit_a_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    exhibit_a_title.paragraph_format.space_after = Pt(12)
    add_text(exhibit_a_title, "PROJECT SCOPE", bold=True)

    add_heading(doc, "A. Included Initial Design and Development")
    add_body(
        doc,
        "WD will design and develop a modern, mobile-friendly responsive business website for Austin Surface Pros. \u201cMobile-friendly responsive\u201d means the layout will adapt for common phone, tablet, laptop, and desktop browser widths. The included website will contain:",
    )
    for item in [
        "Home page with core value proposition, calls to action, and summary of primary services.",
        "Services overview and service-detail content covering agreed offerings, including asphalt and seal coating, parking-lot striping and markings (including fire lanes and hash marks), concrete and asphalt repair, surface or steel coatings, signage, wheel stops, speed bumps, and related commercial safety improvements.",
        "About page presenting Client's business story and experience, based on Client-supplied and Client-approved facts.",
        "Gallery for Client-supplied project photographs and media.",
        "Resources page and Blog with administrator tools for creating, editing, saving drafts, publishing, and deleting posts, including text formatting, links, thumbnail images, and image uploads.",
        "Contact / Estimate Request form that sends a notification email to the recipient selected by Client, initially austinsurfacepros@gmail.com. Final form fields and collection notices are subject to Client approval.",
        "Privacy Policy and Terms and Conditions pages configured to reflect the known website functionality, subject to Client's legal review and continuing approval.",
        "YouTube video embeds for Client-supplied video links. Video hosting, transcoding, channel management, advertising controls, and YouTube's independent practices are excluded.",
        "Administrator-only authentication for Wayne Filer and individual Client employees authorized by Client. Public registration, public customer accounts, public comments, and topic subscriptions are excluded.",
        "Basic on-page search optimization for included public pages, including descriptive titles, metadata, headings, internal links, and reasonable keyword research. Search placement and results are not guaranteed.",
        "Initial Google Business Profile creation or setup assistance. Client or its designated employee is responsible for ongoing posts, messages, reviews, responses, business-hour changes, verification requests, and platform compliance after setup.",
        "Domain transfer and configuration, DNS, TLS certificate configuration, content delivery, production deployment, and AWS cloud hosting administration as described in Exhibit B.",
    ]:
        add_bullet(doc, item, numpr)

    add_heading(doc, "B. Design Review and Content")
    add_body(
        doc,
        "The Initial Fee includes the base design, implementation of the included pages and features, population of Client Materials available before launch, and two consolidated design-feedback rounds. Client will supply final logo files, photographs, video links, service copy, biography or company story, public business information, and any required legal or licensing details. WD may reasonably resize, crop, compress, or format supplied media for web delivery but is not responsible for substantive media production unless separately approved at the Additional Services Rate.",
    )

    add_heading(doc, "C. Expressly Excluded or Future Work")
    add_body(doc, "The following are outside the Initial Fee and Monthly Fee unless added by signed change order or separate agreement:")
    for item in [
        "QuickBooks integration or replacement, automated invoicing, online payment processing, customer portals, or accounting workflows.",
        "Satellite imagery, property measurement, takeoff, estimating, mapping, geospatial analysis, or automated quote-generation tools.",
        "Custom mobile applications, public user accounts, public comments, subscriptions, forums, memberships, e-commerce, or advanced customer relationship management.",
        "Original photography, video production, custom illustration, logo redesign, copywriting beyond light formatting, or legal drafting and legal review.",
        "Paid advertising, recurring social-media management, guaranteed search optimization results, reputation management, or ongoing Google Business Profile management.",
        "Business-domain mailbox subscriptions, email migration, employee mailbox administration, or managed IT services.",
        "Any page, feature, integration, workflow, or data use not expressly listed as included above.",
    ]:
        add_bullet(doc, item, numpr)

    add_page_break(doc)
    exhibit_b = doc.add_paragraph()
    exhibit_b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    exhibit_b.paragraph_format.space_after = Pt(3)
    add_text(exhibit_b, "EXHIBIT B", bold=True)
    exhibit_b_title = doc.add_paragraph()
    exhibit_b_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    exhibit_b_title.paragraph_format.space_after = Pt(12)
    add_text(exhibit_b_title, "MANAGED HOSTING AND SUPPORT SCHEDULE", bold=True)

    add_heading(doc, "A. Included Managed Services")
    for item in [
        "Operation of the production website using AWS services selected by WD for the agreed architecture.",
        "Domain, DNS, TLS certificate, and content-delivery configuration and ordinary renewals, subject to the operating-cost terms below.",
        "Reasonable monitoring, routine maintenance, security updates within WD's control, deployment support, and backups appropriate to the selected architecture.",
        "Included content edits and meeting time described in Section 5 of the Agreement.",
        "Reasonable troubleshooting of the website and WD-managed configuration. Third-party vendor support, material reconfiguration, and new development may be Additional Services.",
    ]:
        add_bullet(doc, item, numpr)

    add_heading(doc, "B. Operating-Cost Allowance and Overage")
    add_body(
        doc,
        "The $50.00 Monthly Fee includes up to $10.00 per calendar month of aggregate third-party operating costs attributable to the website. These costs may include AWS compute, database, object or file storage, bandwidth or data transfer, content delivery, DNS, domain registration and renewal (amortized monthly), certificate costs that are not free, logging, monitoring, backup, and transactional email. WD may reasonably allocate shared-provider charges based on actual metering, provider tags, account records, or a reasonable proportional method.",
    )
    add_body(
        doc,
        "If attributable operating costs exceed $10.00 in a calendar month, Client will pay the actual excess plus a 15% administration and handling margin. WD will include reasonable supporting detail on the invoice and will warn Client in advance when practicable. The Parties acknowledge that usage cost does not always correlate directly with visitor count, page views, storage gigabytes, or a single metric; provider invoices and reasonable allocation records control.",
    )
    add_body(
        doc,
        "If attributable operating costs exceed $10.00 for two consecutive months or materially increase, WD may propose a different hosting architecture, usage limit, Monthly Fee, or direct-billing arrangement. No replacement pricing is effective without written approval, but the existing actual-excess-plus-15% overage continues while the current Agreement remains in effect.",
    )
    add_body(
        doc,
        "Optional business mailbox subscriptions, paid advertising, paid media storage or streaming selected by Client, payment-processing fees, premium plug-ins or software licenses, and other services expressly excluded in the Agreement are not part of the $10.00 allowance and are billed or paid as separately approved.",
    )

    add_heading(doc, "C. Support Requests and Response Targets")
    add_body(
        doc,
        "Client may submit support requests at any time using the contact method designated by WD. Meetings are scheduled by mutual availability; WD does not maintain guaranteed office hours under this Agreement. WD will use commercially reasonable efforts to acknowledge a confirmed production outage or material security issue within one business day, an impaired but usable feature within two business days, and a routine question or minor request within three business days. These are response targets, not resolution guarantees or service credits.",
    )
    add_body(
        doc,
        "This Schedule does not include 24/7 staffing, continuous human monitoring, guaranteed uptime, guaranteed recovery time, guaranteed data restoration, or service credits. WD may make reasonable exceptions or provide faster assistance without creating an ongoing obligation.",
    )

    add_heading(doc, "D. Maintenance and Service Exclusions")
    add_body(
        doc,
        "Availability and response calculations, if discussed, exclude planned or emergency maintenance, AWS or other third-party failures, Internet or utility failures, domain or registry events, Client acts or omissions, Client-requested changes, compromised Client credentials, traffic attacks, illegal use, force-majeure events, and suspension for nonpayment or security protection. WD may temporarily restrict traffic or functionality when reasonably necessary to protect the website, infrastructure, WD, Client, or third parties.",
    )

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.text:
                set_run_font(run, run.font.size or BODY_SIZE, run.bold, run.italic)

    doc.core_properties.title = "Austin Surface Pros Web Development Agreement"
    doc.core_properties.subject = "Web development, managed hosting, and support agreement"
    doc.core_properties.author = "WD Web Solutions LLC"
    doc.core_properties.last_modified_by = "WD Web Solutions LLC"
    doc.core_properties.comments = "Client review copy; revised provisions are highlighted yellow."

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))


if __name__ == "__main__":
    build()
